"""DOCX document parser using python-docx."""

from pathlib import Path
from typing import Union

from app.file_processing.exceptions import (
    DocumentParsingError,
    ParserNotAvailableError,
)
from app.file_processing.models.extracted_document import ExtractedDocument
from .base import DocumentParser


class DOCXParser(DocumentParser):
    """Parser for Microsoft Word DOCX files."""

    SUPPORTED_EXTENSIONS = frozenset({".docx"})

    @classmethod
    def parse(cls, file_path: Union[str, Path]) -> ExtractedDocument:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        try:
            import docx
        except ImportError as exc:
            raise ParserNotAvailableError(
                "python-docx is not installed."
            ) from exc

        try:
            document = docx.Document(path)
        except Exception as exc:
            raise DocumentParsingError(
                f"Failed to read DOCX file: {exc}",
                source_type="DOCX",
            ) from exc

        # Extract text from all paragraphs
        text_parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        full_text = "\n".join(text_parts)

        # Normalize text
        full_text = full_text.replace("\x00", "")
        full_text = full_text.replace("\r\n", "\n").replace("\r", "\n")
        full_text = full_text.strip()

        if not full_text:
            raise DocumentParsingError(
                "DOCX contains no extractable text.",
                source_type="DOCX",
            )

        # Extract metadata
        core_properties = document.core_properties

        title = core_properties.title or None
        author = core_properties.author or None

        # Fallback to the first Title paragraph if document metadata has no title
        if not title:
            for paragraph in document.paragraphs:
                if paragraph.style and paragraph.style.name == "Title":
                    title = paragraph.text.strip()
                    if title:
                        break

        return ExtractedDocument(
            text=full_text,
            filename=path.name,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            page_count=None,
            title=title,
            author=author,
            source_type="DOCX",
        )