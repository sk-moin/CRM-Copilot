"""
app/rag/loaders/parser.py

Document parser used by the RAG ingestion pipeline.

Supported formats
-----------------
- PDF (.pdf)
- Microsoft Word (.docx)
- Plain Text (.txt)
- Markdown (.md)

This module is responsible only for extracting text from files.
Chunking, embeddings, and indexing are handled elsewhere.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

from app.rag.exceptions import (
    DocumentParsingError,
    UnsupportedDocumentTypeError,
)
from app.rag.models.extracted_document import ExtractedDocument


class DocumentParser:
    """Extract text from supported document formats."""

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
        ".md",
    }

    def parse(
        self,
        file_path: str | Path,
    ) -> ExtractedDocument:
        """
        Parse a document into plain text.

        Args:
            file_path:
                Path to the document.

        Returns:
            ExtractedDocument

        Raises:
            UnsupportedDocumentTypeError
            DocumentParsingError
        """

        path = Path(file_path)

        if not path.exists():
            raise DocumentParsingError(
                f"Document not found: {path}"
            )

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentTypeError(
                f"Unsupported document type: {extension}"
            )

        try:
            if extension == ".pdf":
                text = self._parse_pdf(path)

            elif extension == ".docx":
                text = self._parse_docx(path)

            else:
                text = path.read_text(
                    encoding="utf-8",
                    errors="ignore",
                )

            return ExtractedDocument(
                content=text.strip(),
                filename=path.name,
                mime_type=self._mime_type(extension),
            )

        except Exception as exc:
            raise DocumentParsingError(
                f"Failed to parse '{path.name}'"
            ) from exc

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        """Extract text from PDF."""

        reader = PdfReader(path)

        pages = []

        for page in reader.pages:
            pages.append(page.extract_text() or "")

        return "\n".join(pages)

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """Extract text from DOCX."""

        document = Document(path)

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        return "\n".join(paragraphs)

    @staticmethod
    def _mime_type(extension: str) -> str:
        """Return MIME type for supported extensions."""

        mapping = {
            ".pdf": "application/pdf",
            ".docx": (
                "application/"
                "vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            ".txt": "text/plain",
            ".md": "text/markdown",
        }

        return mapping.get(
            extension,
            "application/octet-stream",
        )


def create_document_parser() -> DocumentParser:
    """Factory used for dependency injection."""

    return DocumentParser()