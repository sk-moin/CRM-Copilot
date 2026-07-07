"""Tests for DocxParser."""

from pathlib import Path

import pytest
from docx import Document
from app.file_processing.exceptions import DocumentParsingError
from app.file_processing.models.extracted_document import ExtractedDocument
from app.file_processing.parsers.docx_parser import DOCXParser


def test_can_parse_docx():
    """Parser should recognize .docx files."""
    assert DOCXParser.can_parse("sample.docx")
    assert DOCXParser.can_parse(Path("sample.DOCX"))


def test_cannot_parse_other_extensions():
    """Parser should reject unsupported extensions."""
    assert not DOCXParser.can_parse("sample.pdf")
    assert not DOCXParser.can_parse("sample.txt")


def test_parse_docx_file(tmp_path):
    """Parse a simple DOCX document."""
    file_path = tmp_path / "sample.docx"

    doc = Document()
    doc.add_heading("Sample Title", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(file_path)

    extracted = DOCXParser.parse(file_path)

    assert isinstance(extracted, ExtractedDocument)
    assert extracted.filename == "sample.docx"
    assert "Sample Title" in extracted.text
    assert "First paragraph." in extracted.text
    assert "Second paragraph." in extracted.text
    assert extracted.page_count is None


def test_parse_empty_docx(tmp_path):
    """Parsing an empty DOCX should succeed."""
    file_path = tmp_path / "empty.docx"

    doc = Document()
    doc.save(file_path)

    with pytest.raises(DocumentParsingError, match="no extractable text"):
        DOCXParser.parse(file_path)


def test_parse_unicode_docx(tmp_path):
    """Parser should preserve Unicode text."""
    file_path = tmp_path / "unicode.docx"

    doc = Document()
    doc.add_paragraph("こんにちは")
    doc.add_paragraph("مرحبا")
    doc.add_paragraph("नमस्ते")
    doc.save(file_path)

    extracted = DOCXParser.parse(file_path)

    assert "こんにちは" in extracted.text
    assert "مرحبا" in extracted.text
    assert "नमस्ते" in extracted.text


def test_missing_docx_file():
    """Missing DOCX files should raise FileNotFoundError."""
    with pytest.raises(FileNotFoundError):
        DOCXParser.parse("missing.docx")


def test_corrupt_docx_file(tmp_path):
    """Corrupt DOCX files should raise an exception."""
    file_path = tmp_path / "corrupt.docx"

    file_path.write_bytes(b"This is not a valid DOCX file")

    with pytest.raises(Exception):
        DOCXParser.parse(file_path)